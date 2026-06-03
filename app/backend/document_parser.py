"""Document parser module for extracting text from PDF and DOCX startup pitch files."""

import io
from pypdf import PdfReader
import docx
from app.backend.exceptions import DocumentParsingError

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes.

    Args:
        file_bytes: The raw bytes of the PDF file.

    Returns:
        The extracted and cleaned text content.

    Raises:
        DocumentParsingError: If PDF parsing fails.
    """
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text_list = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_list.append(page_text)

        full_text = "\n".join(text_list)
        if not full_text.strip():
            raise DocumentParsingError("PDF file appears to be empty or contains only images/scans.")

        return clean_text(full_text)
    except Exception as e:
        if isinstance(e, DocumentParsingError):
            raise e
        raise DocumentParsingError(f"Failed to parse PDF document: {str(e)}") from e

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes.

    Args:
        file_bytes: The raw bytes of the DOCX file.

    Returns:
        The extracted and cleaned text content.

    Raises:
        DocumentParsingError: If DOCX parsing fails.
    """
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text_list = []

        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_list.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_list.append(cell.text)

        full_text = "\n".join(text_list)
        if not full_text.strip():
            raise DocumentParsingError("DOCX file appears to be empty.")

        return clean_text(full_text)
    except Exception as e:
        raise DocumentParsingError(f"Failed to parse DOCX document: {str(e)}") from e

def clean_text(text: str) -> str:
    """Perform basic text sanitization (whitespace reduction, normalization)."""
    # Replace multiple newlines and spaces
    lines = [line.strip() for line in text.split("\n")]
    non_empty_lines = [line for line in lines if line]
    cleaned = "\n".join(non_empty_lines)
    # Replace multiple internal spaces
    return " ".join(cleaned.split())

def extract_text(file_bytes: bytes, file_name: str) -> str:
    """Parse document based on filename extension.

    Args:
        file_bytes: The raw bytes of the file.
        file_name: The name of the file (to determine type).

    Returns:
        Extracted and sanitized text content.

    Raises:
        DocumentParsingError: For unsupported types or parsing failure.
    """
    ext = file_name.split(".")[-1].lower()
    if ext == "pdf":
        return parse_pdf(file_bytes)
    if ext in ["docx", "doc"]:
        return parse_docx(file_bytes)
    raise DocumentParsingError(f"Unsupported file format: .{ext}. Please upload a PDF or DOCX file.")
